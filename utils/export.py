"""Export helpers for PDF, DOCX, and Excel."""

import io
from datetime import datetime


def _safe_text(text: str) -> str:
    """Sanitize text for fpdf2 — replace characters that can't be rendered in Helvetica."""
    if not text:
        return ""
    # Replace common problematic characters
    replacements = {
        "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"',
        "\u2013": "-", "\u2014": "--", "\u2026": "...", "\u00a0": " ",
        "\u2022": "*", "\u2023": ">", "\u25cf": "*", "\u25cb": "o",
        "\u2010": "-", "\u2011": "-", "\u2012": "-",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # Remove any remaining non-latin1 characters
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _line(pdf, text: str = "", h: int = 6):
    """Write one full-width line and return the cursor to the left margin.

    multi_cell defaults to new_x=RIGHT, which parks the cursor at the right
    margin and leaves the next full-width (w=0) call with zero usable width.
    """
    pdf.cell(w=0, h=h, text=_safe_text(text), new_x="LMARGIN", new_y="NEXT")


def _safe_multi_cell(pdf, text: str, h: int = 5):
    """Write wrapped text to PDF, catching any rendering errors gracefully."""
    text = _safe_text(text)
    if not text.strip():
        pdf.ln(h)
        return
    try:
        pdf.multi_cell(w=0, h=h, text=text, new_x="LMARGIN", new_y="NEXT")
    except Exception:
        # Fallback: write line by line, skip any that fail
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            try:
                _line(pdf, line[:120], h)
            except Exception:
                pass


def export_contract_pdf(contract_text: str, metadata: dict = None) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    pdf.set_font("Helvetica", "B", 16)
    title = metadata.get("contract_type", "Contract") if metadata else "Contract"
    pdf.cell(0, 10, _safe_text(title), new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    if metadata:
        pdf.set_font("Helvetica", "", 10)
        _line(pdf, f"Date: {metadata.get('effective_date', 'N/A')}")
        _line(pdf, f"Status: {metadata.get('status', 'Draft')}")
        pdf.ln(5)

    pdf.set_font("Helvetica", "", 11)
    for line in (contract_text or "").split("\n"):
        _safe_multi_cell(pdf, line, h=6)

    return bytes(pdf.output())


def export_contract_docx(contract_text: str, metadata: dict = None) -> bytes:
    from docx import Document

    doc = Document()
    title = metadata.get("contract_type", "Contract") if metadata else "Contract"
    doc.add_heading(title, level=0)

    if metadata:
        doc.add_paragraph(f"Date: {metadata.get('effective_date', 'N/A')}")
        doc.add_paragraph(f"Status: {metadata.get('status', 'Draft')}")

    for para in contract_text.split("\n\n"):
        doc.add_paragraph(para)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_redline_docx(playbook: dict, contract_type: str = "Contract") -> bytes:
    """Negotiation playbook as a Word document: strike-through original, bold redline."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading(f"Negotiation Playbook - {contract_type}", level=0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    summary = playbook.get("negotiation_summary", "")
    if summary:
        doc.add_heading("Strategy", level=1)
        doc.add_paragraph(summary)

    must_win = playbook.get("must_win") or []
    tradeable = playbook.get("tradeable") or []
    if must_win or tradeable:
        doc.add_heading("Positions at a glance", level=1)
        if must_win:
            doc.add_paragraph("Must win:")
            for item in must_win:
                doc.add_paragraph(str(item), style="List Bullet")
        if tradeable:
            doc.add_paragraph("Tradeable:")
            for item in tradeable:
                doc.add_paragraph(str(item), style="List Bullet")

    doc.add_heading("Clause-by-clause redlines", level=1)
    for position in playbook.get("positions", []):
        doc.add_heading(f"[{position.get('severity', 'Medium')}] "
                        f"{position.get('risk_type', 'Clause')}", level=2)
        doc.add_paragraph(f"Priority: {position.get('priority', '')}    "
                          f"Leverage: {position.get('leverage', 'Medium')}")

        original = position.get("original_text", "")
        if original:
            doc.add_paragraph("Current language:")
            run = doc.add_paragraph().add_run(original)
            run.font.strike = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        redline = position.get("redlined_text", "")
        if redline:
            doc.add_paragraph("Proposed language:")
            run = doc.add_paragraph().add_run(redline)
            run.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x5A, 0x9C)
            run.font.size = Pt(11)

        rationale = position.get("rationale", "")
        if rationale:
            doc.add_paragraph(f"Why: {rationale}")

        ladder = position.get("fallback_ladder") or {}
        if ladder:
            doc.add_paragraph("Fallback ladder:")
            for key, label in (("ideal", "Ideal"), ("acceptable", "Acceptable"),
                               ("walk_away", "Walk away")):
                if ladder.get(key):
                    doc.add_paragraph(f"{label}: {ladder[key]}", style="List Bullet")

        if position.get("counterparty_objection"):
            doc.add_paragraph(f"Expected objection: {position['counterparty_objection']}")
        if position.get("response"):
            doc.add_paragraph(f"Our response: {position['response']}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_risk_report_pdf(analysis: dict) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Contract Risk Analysis Report", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    pdf.set_font("Helvetica", "", 11)
    _line(pdf, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    _line(pdf, f"Risk Score: {analysis.get('overall_risk_score', 'N/A')} / 100")
    _line(pdf, f"Risk Level: {analysis.get('risk_level', 'N/A')}")
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 13)
    _line(pdf, "Executive Summary", h=8)
    pdf.set_font("Helvetica", "", 10)
    _safe_multi_cell(pdf, analysis.get("executive_summary", "N/A"), h=6)
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 13)
    _line(pdf, "Risky Clauses", h=8)
    pdf.ln(2)

    for clause in analysis.get("risky_clauses", []):
        severity = _safe_text(clause.get("severity", "N/A"))
        risk_type = _safe_text(clause.get("risk_type", "Risk"))
        explanation = _safe_text(clause.get("explanation", ""))
        recommendation = _safe_text(clause.get("recommendation", ""))

        pdf.set_font("Helvetica", "B", 10)
        _line(pdf, f"[{severity}] {risk_type}"[:80])

        if explanation:
            pdf.set_font("Helvetica", "", 9)
            _safe_multi_cell(pdf, explanation)

        if recommendation:
            pdf.set_font("Helvetica", "", 9)
            _safe_multi_cell(pdf, "Recommendation: " + recommendation)

        pdf.ln(3)

    # Missing protections
    missing = analysis.get("missing_protections", [])
    if missing:
        pdf.set_font("Helvetica", "B", 13)
        _line(pdf, "Missing Protections", h=8)
        pdf.set_font("Helvetica", "", 9)
        for m in missing:
            text = _safe_text(f"- {m.get('protection', '')}: {m.get('recommendation', '')}")
            _safe_multi_cell(pdf, text)
        pdf.ln(3)

    # Negotiation points
    negotiations = analysis.get("negotiation_points", [])
    if negotiations:
        pdf.set_font("Helvetica", "B", 13)
        _line(pdf, "Negotiation Points", h=8)
        pdf.set_font("Helvetica", "", 9)
        for n in negotiations:
            text = _safe_text(f"[{n.get('priority', '')}] {n.get('point', '')} - {n.get('suggested_change', '')}")
            _safe_multi_cell(pdf, text)
        pdf.ln(3)

    return bytes(pdf.output())


def export_contracts_excel(df) -> bytes:
    with io.BytesIO() as buf:
        df.to_excel(buf, index=False, engine="openpyxl")
        return buf.getvalue()
